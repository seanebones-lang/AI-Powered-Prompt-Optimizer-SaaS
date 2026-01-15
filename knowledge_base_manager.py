"""
Knowledge Base Manager
Upload and manage custom domain knowledge for enhanced prompt optimization.

Features:
- Document upload (PDF, TXT, MD, DOCX)
- Automatic chunking and embedding
- Semantic search
- Private vector stores per user
- Knowledge base versioning
"""

import logging
import hashlib
import os
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
import json

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents an uploaded document."""
    id: str
    filename: str
    file_type: str
    content: str
    file_size: int
    content_hash: str
    chunks: List[str]
    metadata: Dict[str, Any]
    uploaded_at: str


@dataclass
class SearchResult:
    """Result from knowledge base search."""
    chunk: str
    document_id: str
    document_name: str
    relevance_score: float
    metadata: Dict[str, Any]


class KnowledgeBaseManager:
    """Manages custom knowledge bases for users."""
    
    def __init__(self, storage_path: str = "./knowledge_bases"):
        self.logger = logging.getLogger(__name__)
        self.storage_path = storage_path
        self._ensure_storage_exists()
    
    def _ensure_storage_exists(self):
        """Ensure storage directory exists."""
        os.makedirs(self.storage_path, exist_ok=True)
    
    def create_knowledge_base(
        self,
        user_id: int,
        name: str,
        description: Optional[str] = None,
        domain: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Create a new knowledge base for a user.
        
        Args:
            user_id: User ID
            name: Knowledge base name
            description: Optional description
            domain: Optional domain/industry
            
        Returns:
            Knowledge base info
        """
        from database import db
        
        kb = db.create_knowledge_base(
            user_id=user_id,
            name=name,
            description=description,
            domain=domain,
            is_private=True
        )
        
        if kb:
            # Create storage directory for this KB
            kb_path = os.path.join(self.storage_path, f"kb_{kb.id}")
            os.makedirs(kb_path, exist_ok=True)
            
            return {
                "id": kb.id,
                "name": kb.name,
                "description": kb.description,
                "domain": kb.domain,
                "path": kb_path,
                "created_at": kb.created_at.isoformat()
            }
        
        return None
    
    def upload_document(
        self,
        kb_id: int,
        file_path: str,
        filename: str,
        file_type: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Optional[Document]:
        """
        Upload a document to a knowledge base.
        
        Args:
            kb_id: Knowledge base ID
            file_path: Path to the file
            filename: Original filename
            file_type: File type (pdf, txt, md, docx)
            metadata: Optional metadata
            
        Returns:
            Document object
        """
        try:
            # Read file content
            content = self._read_file(file_path, file_type)
            
            if not content:
                self.logger.error(f"Could not read file: {filename}")
                return None
            
            # Calculate hash for deduplication
            content_hash = hashlib.sha256(content.encode()).hexdigest()
            
            # Check for duplicates
            # TODO: Query database for existing documents with same hash
            
            # Chunk the content
            chunks = self._chunk_content(content, file_type)
            
            # Create document object
            doc = Document(
                id=f"doc_{kb_id}_{int(datetime.now().timestamp())}",
                filename=filename,
                file_type=file_type,
                content=content,
                file_size=len(content),
                content_hash=content_hash,
                chunks=chunks,
                metadata=metadata or {},
                uploaded_at=datetime.now().isoformat()
            )
            
            # Save to storage
            self._save_document(kb_id, doc)
            
            # TODO: Generate embeddings for chunks
            # self._generate_embeddings(doc)
            
            self.logger.info(f"Uploaded document: {filename} ({len(chunks)} chunks)")
            
            return doc
        except Exception as e:
            self.logger.error(f"Error uploading document: {str(e)}")
            return None
    
    def _read_file(self, file_path: str, file_type: str) -> Optional[str]:
        """Read content from file based on type."""
        try:
            if file_type in ["txt", "md"]:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif file_type == "pdf":
                # Try to import PDF library
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text = []
                        for page in reader.pages:
                            text.append(page.extract_text())
                        return "\n".join(text)
                except ImportError:
                    self.logger.warning("PyPDF2 not installed, cannot read PDF")
                    return None
            
            elif file_type == "docx":
                # Try to import DOCX library
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = []
                    for para in doc.paragraphs:
                        text.append(para.text)
                    return "\n".join(text)
                except ImportError:
                    self.logger.warning("python-docx not installed, cannot read DOCX")
                    return None
            
            else:
                self.logger.warning(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            self.logger.error(f"Error reading file: {str(e)}")
            return None
    
    def _chunk_content(
        self,
        content: str,
        file_type: str,
        chunk_size: int = 1000,
        overlap: int = 200
    ) -> List[str]:
        """
        Chunk content into smaller pieces for embedding.
        
        Args:
            content: Content to chunk
            file_type: Type of file
            chunk_size: Target chunk size in characters
            overlap: Overlap between chunks
            
        Returns:
            List of content chunks
        """
        if not content:
            return []
        
        chunks = []
        
        # For markdown, try to split on headers
        if file_type == "md":
            # Split on markdown headers
            sections = content.split('\n#')
            for section in sections:
                if section.strip():
                    # If section is too large, split further
                    if len(section) > chunk_size:
                        chunks.extend(self._split_text(section, chunk_size, overlap))
                    else:
                        chunks.append(section.strip())
        else:
            # Default: split on paragraphs and sentences
            chunks = self._split_text(content, chunk_size, overlap)
        
        return [c for c in chunks if c.strip()]  # Remove empty chunks
    
    def _split_text(
        self,
        text: str,
        chunk_size: int,
        overlap: int
    ) -> List[str]:
        """Split text into overlapping chunks."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                sentence_end = max(
                    text.rfind('. ', start, end),
                    text.rfind('! ', start, end),
                    text.rfind('? ', start, end),
                    text.rfind('\n\n', start, end)
                )
                
                if sentence_end > start:
                    end = sentence_end + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
    
    def _save_document(self, kb_id: int, doc: Document):
        """Save document to storage."""
        kb_path = os.path.join(self.storage_path, f"kb_{kb_id}")
        doc_path = os.path.join(kb_path, f"{doc.id}.json")
        
        # Save document metadata and chunks
        doc_data = {
            "id": doc.id,
            "filename": doc.filename,
            "file_type": doc.file_type,
            "file_size": doc.file_size,
            "content_hash": doc.content_hash,
            "chunks": doc.chunks,
            "metadata": doc.metadata,
            "uploaded_at": doc.uploaded_at
        }
        
        with open(doc_path, 'w', encoding='utf-8') as f:
            json.dump(doc_data, f, indent=2)
    
    def search(
        self,
        kb_id: int,
        query: str,
        top_k: int = 5,
        min_score: float = 0.5
    ) -> List[SearchResult]:
        """
        Search knowledge base for relevant content.
        
        Args:
            kb_id: Knowledge base ID
            query: Search query
            top_k: Number of results to return
            min_score: Minimum relevance score
            
        Returns:
            List of search results
        """
        # Load all documents in KB
        kb_path = os.path.join(self.storage_path, f"kb_{kb_id}")
        
        if not os.path.exists(kb_path):
            return []
        
        results = []
        
        # Simple keyword-based search (TODO: Replace with semantic search)
        for filename in os.listdir(kb_path):
            if filename.endswith('.json'):
                doc_path = os.path.join(kb_path, filename)
                
                with open(doc_path, 'r', encoding='utf-8') as f:
                    doc_data = json.load(f)
                
                # Search in chunks
                for i, chunk in enumerate(doc_data['chunks']):
                    score = self._calculate_relevance(query, chunk)
                    
                    if score >= min_score:
                        results.append(SearchResult(
                            chunk=chunk,
                            document_id=doc_data['id'],
                            document_name=doc_data['filename'],
                            relevance_score=score,
                            metadata={
                                "chunk_index": i,
                                "file_type": doc_data['file_type']
                            }
                        ))
        
        # Sort by relevance and return top_k
        results.sort(key=lambda x: x.relevance_score, reverse=True)
        return results[:top_k]
    
    def _calculate_relevance(self, query: str, text: str) -> float:
        """
        Calculate relevance score between query and text.
        
        Simple keyword-based scoring. TODO: Replace with semantic similarity.
        """
        query_lower = query.lower()
        text_lower = text.lower()
        
        # Count matching words
        query_words = set(query_lower.split())
        text_words = set(text_lower.split())
        
        matching_words = query_words.intersection(text_words)
        
        if not query_words:
            return 0.0
        
        # Basic score: percentage of query words found
        score = len(matching_words) / len(query_words)
        
        # Bonus for exact phrase match
        if query_lower in text_lower:
            score += 0.3
        
        return min(1.0, score)
    
    def get_context_for_prompt(
        self,
        kb_id: int,
        prompt: str,
        max_chunks: int = 3
    ) -> str:
        """
        Get relevant context from knowledge base for a prompt.
        
        Args:
            kb_id: Knowledge base ID
            prompt: User's prompt
            max_chunks: Maximum number of chunks to include
            
        Returns:
            Formatted context string
        """
        results = self.search(kb_id, prompt, top_k=max_chunks)
        
        if not results:
            return ""
        
        context_parts = ["Relevant information from your knowledge base:\n"]
        
        for i, result in enumerate(results, 1):
            context_parts.append(f"\n[Source {i}: {result.document_name}]")
            context_parts.append(result.chunk)
            context_parts.append("")
        
        return "\n".join(context_parts)
    
    def list_documents(self, kb_id: int) -> List[Dict[str, Any]]:
        """List all documents in a knowledge base."""
        kb_path = os.path.join(self.storage_path, f"kb_{kb_id}")
        
        if not os.path.exists(kb_path):
            return []
        
        documents = []
        
        for filename in os.listdir(kb_path):
            if filename.endswith('.json'):
                doc_path = os.path.join(kb_path, filename)
                
                with open(doc_path, 'r', encoding='utf-8') as f:
                    doc_data = json.load(f)
                
                documents.append({
                    "id": doc_data['id'],
                    "filename": doc_data['filename'],
                    "file_type": doc_data['file_type'],
                    "file_size": doc_data['file_size'],
                    "chunks": len(doc_data['chunks']),
                    "uploaded_at": doc_data['uploaded_at']
                })
        
        return documents
    
    def delete_document(self, kb_id: int, doc_id: str) -> bool:
        """Delete a document from knowledge base."""
        kb_path = os.path.join(self.storage_path, f"kb_{kb_id}")
        doc_path = os.path.join(kb_path, f"{doc_id}.json")
        
        try:
            if os.path.exists(doc_path):
                os.remove(doc_path)
                self.logger.info(f"Deleted document: {doc_id}")
                return True
            return False
        except Exception as e:
            self.logger.error(f"Error deleting document: {str(e)}")
            return False
    
    def get_statistics(self, kb_id: int) -> Dict[str, Any]:
        """Get statistics for a knowledge base."""
        documents = self.list_documents(kb_id)
        
        if not documents:
            return {
                "document_count": 0,
                "total_chunks": 0,
                "total_size": 0,
                "file_types": {}
            }
        
        total_chunks = sum(doc['chunks'] for doc in documents)
        total_size = sum(doc['file_size'] for doc in documents)
        
        # Count by file type
        file_types = {}
        for doc in documents:
            ft = doc['file_type']
            file_types[ft] = file_types.get(ft, 0) + 1
        
        return {
            "document_count": len(documents),
            "total_chunks": total_chunks,
            "total_size": total_size,
            "file_types": file_types,
            "average_chunks_per_doc": total_chunks / len(documents) if documents else 0
        }
    
    def export_knowledge_base(
        self,
        kb_id: int,
        format: str = "json"
    ) -> str:
        """Export entire knowledge base."""
        documents = self.list_documents(kb_id)
        
        if format == "json":
            return json.dumps(documents, indent=2)
        elif format == "markdown":
            lines = [f"# Knowledge Base Export\n"]
            lines.append(f"Total Documents: {len(documents)}\n")
            
            for doc in documents:
                lines.append(f"\n## {doc['filename']}")
                lines.append(f"- Type: {doc['file_type']}")
                lines.append(f"- Chunks: {doc['chunks']}")
                lines.append(f"- Uploaded: {doc['uploaded_at']}")
            
            return "\n".join(lines)
        else:
            return str(documents)


# Convenience functions
def create_kb(user_id: int, name: str, **kwargs) -> Dict[str, Any]:
    """Create a knowledge base."""
    manager = KnowledgeBaseManager()
    return manager.create_knowledge_base(user_id, name, **kwargs)


def upload_doc(kb_id: int, file_path: str, filename: str, file_type: str) -> Optional[Document]:
    """Upload a document."""
    manager = KnowledgeBaseManager()
    return manager.upload_document(kb_id, file_path, filename, file_type)


def search_kb(kb_id: int, query: str, **kwargs) -> List[SearchResult]:
    """Search a knowledge base."""
    manager = KnowledgeBaseManager()
    return manager.search(kb_id, query, **kwargs)
